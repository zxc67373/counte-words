from docx import Document
import re
import matplotlib.pyplot as plt
import numpy as np

document = Document(r"D:\WeChat\WeChat Files\wxid_d36bky38gch822\FileStorage\File\2022-02\分享test.docx")
newdocument = Document()

AABB = re.compile(r"(.)\1(.)\2")
AA = re.compile(r"(.)\1+")
ABAC = re.compile(r"(.).\1.")


def countword(pattern, document, newdocument):
    flag = 0
    for i in range(len(document.paragraphs)):

        result = pattern.findall(document.paragraphs[i].text)
        if len(result) > 0 and pattern == AA:
            for j in result:
                if j in "-,—，谢, ,[0,1,2,3,4,5,6,7,8,9],q,w,e,r,t,y,u,i,o,p,a,s,d,f,g,h,j,k,l,z,x,c,v,b,n,m],C,太，婆，奶":
                    continue
                #   result.remove(j)
                flag += len(result)
                newdocument.add_paragraph(j)
                newdocument.add_paragraph(document.paragraphs[i].text)
                # para = newdocument.add_paragraph(document.paragraphs[i].text)
                print('AA', j, '\t\t\t\t\t在文中:',
                      document.paragraphs[i].text[document.paragraphs[i].text.index(j):])
        elif len(result) > 0 and pattern == AABB:
            flag += 1
            newdocument.add_paragraph(str(result))
            newdocument.add_paragraph(document.paragraphs[i].text)
            # para = newdocument.add_paragraph(document.paragraphs[i].text)
            print('AABB', result, '\t\t\t\t\t在文中:',
                  document.paragraphs[i].text[document.paragraphs[i].text.index(result[0][0]):])
        elif len(result) > 0 and pattern == ABAC:
            for j in result:
                if j in "-,—，谢,[0,1,2,3,4,5,6,7,8,9],q, ,.,w,e,r,t,y,u,i,o,p,a,s,d,f,g,h,j,k,l,z,x,c,v,b,n,m],C,太，婆，奶":
                    continue
                flag += 1
                newdocument.add_paragraph(j)
                newdocument.add_paragraph(document.paragraphs[i].text)
                # para = newdocument.add_paragraph(document.paragraphs[i].text)
                print('ABAC', j, '\t\t\t\t\t在文中:',
                      document.paragraphs[i].text[document.paragraphs[i].text.index(result[0][0]):])

    return newdocument, flag


newdocument, nums_of_aa = countword(AA, document, newdocument)
newdocument, nums_of_aabb = countword(AABB, document, newdocument)
newdocument, nums_of_abac = countword(ABAC, document, newdocument)

# print(nums_of_aa)
newdocument.save(r"1.docx")
img = np.array([nums_of_abac, nums_of_aabb, nums_of_aa])

plt.pie(img,
        labels=['ABAC', 'AABB', 'AA'],
        colors=["#5d8ca8", "#65a479", "#a564c9"],
        explode=(0, 0.1, 0),
        autopct='%.2f%%',)

plt.title("counter word")
plt.show()
